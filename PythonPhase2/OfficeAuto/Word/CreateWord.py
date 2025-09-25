from docx import Document
doc=Document()
doc.add_heading(',prep',3)
doc.add_paragraph('addd')
doc.add_paragraph('add')
doc.add_heading('dsdsdsd')
doc.save('one.docx')