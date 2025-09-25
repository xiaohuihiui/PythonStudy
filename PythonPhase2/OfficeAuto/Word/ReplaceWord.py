from docx import Document

def wordReplace(doc,oldWord,newWord):
 for paragraph in doc.paragraphs:
  for run in paragraph.runs:
   run.text=run.text.replace(oldWord,newWord)
   #for table in paragraph.tables:
     #for row in table.rows:
         #for cell in row.cells:
   # cell.text=cell.text.replace(oldWord,newWord)#

if __name__ == '__main__':
 doc=Document('one.docx')
 wordReplace(doc,'add','acc')
 doc.save('output1.docx')

